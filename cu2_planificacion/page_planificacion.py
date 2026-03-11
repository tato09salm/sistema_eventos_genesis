import streamlit as st
import pandas as pd
from datetime import date
from io import BytesIO
from docx import Document
from fpdf import FPDF
from auth.roles import requiere_rol, check_rol
from cu2_planificacion import model_evento, model_plan_evento, model_requerimiento, model_cotizacion
from cu3_recursos import model_proveedor, model_recurso, model_orden_compra
from shared.utils import format_currency

# ══════════════════════════════════════════════════════════════════
#  CSS SCOPED — Solo afecta al div#planificacion-page
#  El nav/sidebar queda completamente intacto.
# ══════════════════════════════════════════════════════════════════

THEME_CSS = """
<style>
#planificacion-page {
    --pp-black:  #000000;
    --pp-silver: #c0c0c0;
    --pp-gray:   #808080;
    --pp-white:  #ffffff;
    --pp-maroon: #800000;
    --pp-red:    #ff0000;
    --pp-lime:   #00ff00;
    --pp-blue:   #0000ff;
    --pp-bg:     #0d0d0d;
    --pp-bg2:    #141414;
    --pp-bg3:    #1c1c1c;
    background-color: var(--pp-bg) !important;
    padding: 1.2rem 1.4rem !important;
    border-radius: 10px !important;
    border: 1px solid var(--pp-maroon) !important;
    color: var(--pp-silver) !important;
    font-family: 'Courier New', monospace;
}
#planificacion-page h1 {
    color: var(--pp-white) !important;
    border-bottom: 3px solid var(--pp-maroon) !important;
    padding-bottom: .4rem !important;
    letter-spacing: 3px !important;
    text-transform: uppercase !important;
}
#planificacion-page h2 {
    color: var(--pp-silver) !important;
    border-left: 4px solid var(--pp-maroon) !important;
    padding-left: .5rem !important;
}
#planificacion-page h3 { color: var(--pp-silver) !important; }
#planificacion-page .stTabs [data-baseweb="tab-list"] {
    background-color: var(--pp-bg2) !important;
    border-bottom: 2px solid var(--pp-maroon) !important;
    gap: 2px !important;
}
#planificacion-page .stTabs [data-baseweb="tab"] {
    background-color: var(--pp-bg3) !important;
    color: var(--pp-gray) !important;
    border: 1px solid #333 !important;
    border-bottom: none !important;
    border-radius: 6px 6px 0 0 !important;
    font-weight: 600 !important;
    transition: all .2s ease !important;
}
#planificacion-page .stTabs [aria-selected="true"] {
    background-color: var(--pp-maroon) !important;
    color: var(--pp-white) !important;
    border-color: var(--pp-maroon) !important;
}
#planificacion-page .stTabs [data-baseweb="tab"]:hover {
    background-color: #3a0000 !important;
    color: var(--pp-white) !important;
}
#planificacion-page .stExpander {
    background-color: var(--pp-bg2) !important;
    border: 1px solid var(--pp-gray) !important;
    border-radius: 6px !important;
}
#planificacion-page [data-testid="stExpanderDetails"] {
    background-color: var(--pp-bg2) !important;
}
#planificacion-page [data-testid="stForm"] {
    background-color: var(--pp-bg2) !important;
    border: 1px solid var(--pp-maroon) !important;
    border-radius: 8px !important;
    padding: 1rem !important;
}
#planificacion-page .stTextInput > div > div > input,
#planificacion-page .stNumberInput > div > div > input,
#planificacion-page .stTextArea > div > div > textarea,
#planificacion-page .stDateInput > div > div > input {
    background-color: var(--pp-bg3) !important;
    color: var(--pp-white) !important;
    border: 1px solid var(--pp-gray) !important;
    border-radius: 4px !important;
}
#planificacion-page .stSelectbox > div > div {
    background-color: var(--pp-bg3) !important;
    color: var(--pp-white) !important;
    border: 1px solid var(--pp-gray) !important;
}
#planificacion-page label {
    color: var(--pp-silver) !important;
    font-weight: 600 !important;
    font-size: .85rem !important;
}
#planificacion-page .stButton > button[kind="primary"],
#planificacion-page .stFormSubmitButton > button {
    background-color: var(--pp-maroon) !important;
    color: var(--pp-white) !important;
    border: 1px solid var(--pp-red) !important;
    font-weight: 700 !important;
    transition: all .2s ease !important;
}
#planificacion-page .stButton > button[kind="primary"]:hover,
#planificacion-page .stFormSubmitButton > button:hover {
    background-color: var(--pp-red) !important;
    box-shadow: 0 4px 12px rgba(255,0,0,.3) !important;
}
#planificacion-page .stButton > button {
    background-color: var(--pp-bg3) !important;
    color: var(--pp-silver) !important;
    border: 1px solid var(--pp-gray) !important;
    border-radius: 4px !important;
    transition: all .15s ease !important;
}
#planificacion-page .stButton > button:hover {
    background-color: #2a2a2a !important;
    color: var(--pp-white) !important;
}
#planificacion-page .stButton > button[title="Eliminar"] {
    border-color: var(--pp-red) !important;
    color: var(--pp-red) !important;
}
#planificacion-page .stButton > button[title="Aceptar"] {
    border-color: var(--pp-lime) !important;
    color: var(--pp-lime) !important;
}
#planificacion-page .stDownloadButton > button {
    background-color: #00008b !important;
    color: var(--pp-white) !important;
    border: 1px solid var(--pp-blue) !important;
    font-weight: 700 !important;
    transition: all .2s ease !important;
}
#planificacion-page .stDownloadButton > button:hover {
    background-color: var(--pp-blue) !important;
    box-shadow: 0 4px 12px rgba(0,0,255,.35) !important;
}
#planificacion-page [data-testid="metric-container"] {
    background-color: var(--pp-bg2) !important;
    border: 1px solid var(--pp-maroon) !important;
    border-radius: 8px !important;
    padding: .8rem !important;
}
#planificacion-page [data-testid="metric-container"] label {
    color: var(--pp-gray) !important;
    font-size: .72rem !important;
    letter-spacing: 1px !important;
    text-transform: uppercase !important;
}
#planificacion-page [data-testid="stMetricValue"] {
    color: var(--pp-white) !important;
    font-weight: 700 !important;
}
#planificacion-page .stSuccess > div {
    background-color: #001a00 !important;
    border-left: 4px solid var(--pp-lime) !important;
    color: var(--pp-lime) !important;
}
#planificacion-page .stWarning > div {
    background-color: #1a0f00 !important;
    border-left: 4px solid #ff8800 !important;
    color: #ff8800 !important;
}
#planificacion-page .stError > div {
    background-color: #1a0000 !important;
    border-left: 4px solid var(--pp-red) !important;
    color: var(--pp-red) !important;
}
#planificacion-page .stInfo > div {
    background-color: #00001a !important;
    border-left: 4px solid var(--pp-blue) !important;
    color: #4488ff !important;
}
#planificacion-page p,
#planificacion-page span,
#planificacion-page div  { color: var(--pp-silver) !important; }
#planificacion-page strong,
#planificacion-page b    { color: var(--pp-white) !important; }
#planificacion-page hr   { border-color: var(--pp-maroon) !important; }
#planificacion-page ::-webkit-scrollbar { width: 5px; height: 5px; }
#planificacion-page ::-webkit-scrollbar-track { background: var(--pp-bg); }
#planificacion-page ::-webkit-scrollbar-thumb { background: var(--pp-maroon); border-radius: 3px; }
</style>
"""

# ══════════════════════════════════════════════════════════════════
#  HELPERS GENERALES
# ══════════════════════════════════════════════════════════════════

def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        return d
    try:
        return d.strftime("%d-%m-%Y")
    except Exception:
        return str(d)


def _stats_numericos(valores):
    """Estadísticas descriptivas completas."""
    nums = []
    for v in valores or []:
        if v is None:
            continue
        try:
            nums.append(float(v))
        except Exception:
            continue
    if not nums:
        return {
            "count": 0, "sum": 0.0, "mean": 0.0, "median": 0.0,
            "min": 0.0, "max": 0.0, "range": 0.0,
            "std": 0.0, "cv": 0.0, "q1": 0.0, "q3": 0.0, "iqr": 0.0,
        }
    n = len(nums)
    s = sum(nums)
    mn, mx = min(nums), max(nums)
    mean = s / n
    ns   = sorted(nums)
    mid  = n // 2
    median = ns[mid] if n % 2 else (ns[mid - 1] + ns[mid]) / 2
    std  = (sum((x - mean) ** 2 for x in nums) / n) ** 0.5
    cv   = std / mean * 100 if mean != 0 else 0.0
    q1   = ns[n // 4]
    q3   = ns[(3 * n) // 4]
    return {
        "count": n, "sum": s, "mean": mean, "median": median,
        "min": mn, "max": mx, "range": mx - mn,
        "std": std, "cv": cv, "q1": q1, "q3": q3, "iqr": q3 - q1,
    }


def _conteo(valores):
    cnt = {}
    for v in (valores or []):
        k = str(v); cnt[k] = cnt.get(k, 0) + 1
    return cnt


# ══════════════════════════════════════════════════════════════════
#  GRÁFICOS PLOTLY (fondo oscuro, paleta corporativa)
# ══════════════════════════════════════════════════════════════════

_LAYOUT = dict(
    paper_bgcolor="#0d0d0d", plot_bgcolor="#141414",
    font=dict(color="#c0c0c0"),
    margin=dict(l=10, r=10, t=40, b=30),
)


def _fig_dona(datos: dict, titulo: str):
    try:
        import plotly.graph_objects as go
        cats, vals = list(datos.keys()), list(datos.values())
        if not cats: return None
        palette = ["#800000","#ff0000","#c0c0c0","#808080","#ffffff","#3a0000","#ff4444","#606060"]
        fig = go.Figure(go.Pie(
            labels=cats, values=vals, hole=0.55,
            marker=dict(colors=palette[:len(cats)], line=dict(color="#0d0d0d", width=2)),
            textinfo="label+percent",
            textfont=dict(color="#ffffff", size=9),
            hovertemplate="%{label}: %{value} (%{percent})<extra></extra>",
        ))
        fig.update_layout(title=dict(text=titulo, font=dict(color="#fff", size=12), x=0.01),
                          legend=dict(font=dict(color="#c0c0c0"), bgcolor="#141414"),
                          height=260,
                          annotations=[dict(text=f"<b>{sum(vals)}</b>", x=0.5, y=0.5,
                                            font=dict(size=16, color="#ffffff"), showarrow=False)],
                          **_LAYOUT)
        return fig
    except ImportError:
        return None


def _fig_barras_h(datos: dict, titulo: str, eje_y="Cantidad"):
    try:
        import plotly.graph_objects as go
        cats, vals = list(datos.keys()), list(datos.values())
        if not cats: return None
        fig = go.Figure(go.Bar(
            x=vals, y=cats, orientation="h",
            marker=dict(color=vals,
                        colorscale=[[0,"#3a0000"],[0.5,"#800000"],[1,"#ff4444"]],
                        line=dict(color="#c0c0c0", width=0.4)),
            text=[str(v) for v in vals], textposition="outside",
            textfont=dict(color="#c0c0c0", size=9),
        ))
        fig.update_layout(title=dict(text=titulo, font=dict(color="#fff", size=12), x=0.01),
                          xaxis=dict(title=eje_y, gridcolor="#2a2a2a", color="#808080"),
                          yaxis=dict(gridcolor="#2a2a2a", color="#808080", autorange="reversed"),
                          height=max(160, len(cats) * 38), **_LAYOUT)
        return fig
    except ImportError:
        return None


def _fig_histo(valores: list, titulo: str, unidad: str = ""):
    try:
        import plotly.graph_objects as go
        if not valores or len(valores) < 2: return None
        media = sum(valores) / len(valores)
        fig = go.Figure(go.Histogram(
            x=valores, nbinsx=min(10, len(set(valores))),
            marker=dict(color="#800000", line=dict(color="#c0c0c0", width=0.4)),
            opacity=0.85,
        ))
        fig.add_vline(x=media, line_dash="dash", line_color="#ff0000",
                      annotation_text=f"Media: {media:,.1f}{unidad}",
                      annotation_font_color="#ff0000")
        fig.update_layout(title=dict(text=titulo, font=dict(color="#fff", size=12), x=0.01),
                          xaxis=dict(gridcolor="#2a2a2a", color="#808080"),
                          yaxis=dict(title="Frecuencia", gridcolor="#2a2a2a", color="#808080"),
                          height=210, **_LAYOUT)
        return fig
    except ImportError:
        return None


def _fig_linea(fechas, valores, titulo: str, unidad: str = ""):
    try:
        import plotly.graph_objects as go
        pares = sorted([(f, float(v)) for f, v in zip(fechas, valores) if f and v is not None],
                       key=lambda x: x[0])
        if len(pares) < 2: return None
        xs, ys = [p[0] for p in pares], [p[1] for p in pares]
        fig = go.Figure(go.Scatter(
            x=xs, y=ys, mode="lines+markers",
            line=dict(color="#800000", width=2),
            marker=dict(color="#ff0000", size=7, line=dict(color="#fff", width=1)),
            fill="tozeroy", fillcolor="rgba(128,0,0,0.15)",
            hovertemplate=f"%{{x}}: %{{y:,.2f}}{unidad}<extra></extra>",
        ))
        fig.update_layout(title=dict(text=titulo, font=dict(color="#fff", size=12), x=0.01),
                          xaxis=dict(gridcolor="#2a2a2a", color="#808080"),
                          yaxis=dict(gridcolor="#2a2a2a", color="#808080"),
                          height=220, **_LAYOUT)
        return fig
    except ImportError:
        return None


def _show(fig):
    if fig is not None:
        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


def _panel_stats(s: dict, tipo: str = "currency"):
    fmt = (format_currency if tipo == "currency"
           else (lambda x: f"{x:.2f}" if tipo == "float" else str(int(round(x)))))
    r1 = st.columns(5)
    r1[0].metric("Conteo",   str(s["count"]))
    r1[1].metric("Total",    fmt(s["sum"]))
    r1[2].metric("Promedio", fmt(s["mean"]))
    r1[3].metric("Mediana",  fmt(s["median"]))
    r1[4].metric("Rango",    fmt(s["range"]))
    r2 = st.columns(5)
    r2[0].metric("Mínimo",     fmt(s["min"]))
    r2[1].metric("Máximo",     fmt(s["max"]))
    r2[2].metric("Desv. Std",  fmt(s["std"]))
    r2[3].metric("Coef. Var",  f"{s['cv']:.1f}%")
    r2[4].metric("IQR",        fmt(s["iqr"]))


# ══════════════════════════════════════════════════════════════════
#  PDF CORPORATIVO
# ══════════════════════════════════════════════════════════════════

def _safe(s) -> str:
    return str(s).replace("—","-").replace("–","-").replace("•","-")\
                 .replace("\u2019","'").replace("\u201c",'"').replace("\u201d",'"')\
                 .encode("latin-1","replace").decode("latin-1")


class _PDF(FPDF):
    M=(128,0,0); BK=(10,10,10); W=(255,255,255)
    SL=(192,192,192); GR=(100,100,100); RD=(220,50,50)
    D1=(18,18,18); D2=(28,28,28); D3=(38,38,38)

    def __init__(self, titulo="", subtitulo=""):
        super().__init__()
        self._titulo, self._sub = titulo, subtitulo
        self.set_auto_page_break(auto=True, margin=22)

    def header(self):
        self.set_fill_color(*self.M);  self.rect(0,0,210,16,"F")
        self.set_fill_color(*self.BK); self.rect(0,16,210,2,"F")
        self.set_y(2)
        self.set_font("Arial","B",12); self.set_text_color(*self.W)
        self.cell(170,12,_safe(self._titulo),align="C")
        self.set_font("Arial","",8);   self.set_text_color(220,180,180)
        self.cell(0,12,f"Pag. {self.page_no()}",align="R")
        if self._sub:
            self.set_fill_color(50,0,0); self.rect(0,18,210,6,"F")
            self.set_xy(0,18); self.set_font("Arial","I",8)
            self.set_text_color(220,180,180); self.cell(0,6,_safe(self._sub),align="C")
            self.ln(8)
        else:
            self.ln(6)

    def footer(self):
        self.set_y(-13)
        self.set_fill_color(*self.M); self.rect(0,self.get_y(),210,13,"F")
        self.set_font("Arial","I",7.5); self.set_text_color(220,180,180)
        self.cell(0,10,f"Sistema de Gestion de Eventos  |  {date.today().strftime('%d/%m/%Y')}  |  Generado automaticamente",align="C")

    def seccion(self, texto: str):
        self.ln(3)
        self.set_fill_color(*self.M); self.rect(self.l_margin,self.get_y(),190,8,"F")
        self.set_fill_color(*self.RD); self.rect(self.l_margin,self.get_y(),3,8,"F")
        self.set_xy(self.l_margin+5,self.get_y())
        self.set_font("Arial","B",9); self.set_text_color(*self.W)
        self.cell(185,8,_safe(texto.upper()),ln=1)
        self.ln(1); self.set_text_color(*self.SL)

    def kv(self, label: str, valor: str, alt: bool=False):
        bl = self.D2 if alt else self.D1
        bv = self.D3 if alt else self.D2
        y  = self.get_y()
        self.set_fill_color(*bl); self.rect(self.l_margin,y,62,7,"F")
        self.set_font("Arial","B",8); self.set_text_color(*self.GR)
        self.set_xy(self.l_margin+2,y); self.cell(60,7,_safe(label)+":")
        self.set_fill_color(*bv); self.rect(self.l_margin+62,y,128,7,"F")
        self.set_font("Arial","",8); self.set_text_color(*self.W)
        self.set_xy(self.l_margin+64,y); self.cell(126,7,_safe(str(valor)),ln=1)
        self.set_text_color(*self.SL)

    def sep(self):
        self.set_draw_color(*self.M); self.set_line_width(0.4)
        self.line(self.l_margin,self.get_y(),200,self.get_y()); self.ln(3)

    def th(self, cols, widths):
        self.set_fill_color(*self.M); self.set_text_color(*self.W)
        self.set_font("Arial","B",7.5)
        for c,w in zip(cols,widths): self.cell(w,7,_safe(c),fill=True,align="C")
        self.ln()

    def tf(self, vals, widths, idx=0, aligns=None):
        self.set_fill_color(*(self.D1 if idx%2==0 else self.D2))
        self.set_font("Arial","",7.5); self.set_text_color(*self.SL)
        als = aligns or ["L"]*len(vals)
        for v,w,a in zip(vals,widths,als): self.cell(w,6,_safe(str(v)),fill=True,align=a)
        self.ln()

    def tt(self, vals, widths, aligns=None):
        self.set_fill_color(60,0,0); self.set_font("Arial","B",8); self.set_text_color(*self.W)
        als = aligns or ["L"]*len(vals)
        for v,w,a in zip(vals,widths,als): self.cell(w,7,_safe(str(v)),fill=True,align=a)
        self.ln()

    def bloque_stats(self, s: dict, tipo: str="currency"):
        fmt = (format_currency if tipo=="currency"
               else (lambda x: f"{x:.2f}" if tipo=="float" else str(int(round(x)))))
        self.seccion("Estadisticas Descriptivas")
        pares = [
            ("Conteo",           str(s["count"])),
            ("Total",            fmt(s["sum"])),
            ("Promedio",         fmt(s["mean"])),
            ("Mediana",          fmt(s["median"])),
            ("Minimo",           fmt(s["min"])),
            ("Maximo",           fmt(s["max"])),
            ("Rango",            fmt(s["range"])),
            ("Desv. Estandar",   fmt(s["std"])),
            ("Coef. Variacion",  f"{s['cv']:.2f}%"),
            ("Q1  (percentil 25)", fmt(s["q1"])),
            ("Q3  (percentil 75)", fmt(s["q3"])),
            ("IQR (Q3 - Q1)",    fmt(s["iqr"])),
        ]
        for i,(k,v) in enumerate(pares): self.kv(k,v,alt=(i%2==1))
        self.ln(2)

    def bloque_dist(self, cnt: dict, titulo: str):
        if not cnt: return
        self.seccion(titulo)
        total = sum(cnt.values()); mx = max(cnt.values()) if cnt else 1
        self.set_font("Courier","B",7.5); self.set_text_color(*self.GR)
        self.cell(65,6,"Categoria"); self.cell(20,6,"Cant.",align="R")
        self.cell(20,6,"%",align="R"); self.cell(0,6,"Visual (# = unidades)",ln=1)
        self.sep()
        for i,(k,v) in enumerate(sorted(cnt.items(),key=lambda x:-x[1])):
            self.set_fill_color(*(self.D1 if i%2==0 else self.D2))
            y=self.get_y(); self.rect(self.l_margin,y,190,6,"F"); self.set_xy(self.l_margin,y)
            pct=v/total*100; bars="#"*int(v/mx*45)
            self.set_font("Courier","B",7); self.set_text_color(*self.SL)
            self.cell(65,6,_safe(str(k)[:30]))
            self.set_font("Courier","",7)
            self.cell(20,6,str(v),align="R"); self.cell(20,6,f"{pct:.1f}%",align="R")
            self.set_text_color(*self.M); self.cell(0,6,bars[:45],ln=1)
        self.ln(2); self.set_text_color(*self.SL)

    def nota(self, texto: str):
        self.set_font("Arial","I",7.5); self.set_text_color(*self.GR)
        self.multi_cell(0,5,_safe(texto)); self.ln(1); self.set_text_color(*self.SL)


# ── PDF por módulo ────────────────────────────────────────────────

def _bytes(pdf): 
    raw = pdf.output(dest="S")
    return raw.encode("latin-1") if isinstance(raw, str) else bytes(raw)


def _pdf_eventos(rows):
    pdf = _PDF("Reporte de Eventos", f"Total: {len(rows)} eventos")
    pdf.add_page()
    s = _stats_numericos([float(r[5] or 0) for r in rows])
    pdf.bloque_stats(s,"currency")
    pdf.bloque_dist(_conteo([r[6] for r in rows]), "Distribucion por Estado")
    pdf.bloque_dist(_conteo([r[2] for r in rows]), "Distribucion por Tipo")
    pdf.seccion("Listado de Eventos")
    cols=["Nombre","Tipo","Lugar","Fecha","Monto S/","Estado"]
    ws  =[44,24,34,22,28,28]; als=["L","C","L","C","R","C"]
    pdf.th(cols,ws); tot=0.0
    for i,r in enumerate(rows):
        m=float(r[5] or 0); tot+=m
        pdf.tf([str(r[1])[:22],str(r[2]),str(r[3] or "")[:18],
                _fmt_date(r[4]),format_currency(m),str(r[6])],ws,i,als)
    pdf.tt(["TOTAL","","","",format_currency(tot),""],ws,als)
    pdf.sep(); pdf.nota(f"Generado: {date.today().strftime('%d/%m/%Y')}. Montos en Soles (S/).")
    return _bytes(pdf)


def _pdf_planes(planes, nombre):
    pdf = _PDF("Reporte de Planes", f"Evento: {nombre}  |  Total: {len(planes)} planes")
    pdf.add_page()
    s = _stats_numericos([float(p[2] or 0) for p in planes])
    pdf.bloque_stats(s,"currency")
    pdf.bloque_dist(_conteo([str(p[3]) for p in planes]),"Distribucion por Estado")
    pdf.seccion("Listado de Planes")
    cols=["Fecha Elab.","Presupuesto S/","Estado","Descripcion"]
    ws  =[28,34,28,90]; als=["C","R","C","L"]
    pdf.th(cols,ws); tot=0.0
    for i,p in enumerate(planes):
        tot+=float(p[2] or 0)
        pdf.tf([_fmt_date(p[1]),format_currency(p[2] or 0),str(p[3]),str(p[4] or "")[:50]],ws,i,als)
    pdf.tt(["",format_currency(tot),"TOTAL",""],ws,als)
    pdf.sep(); pdf.nota(f"Evento: {nombre}. Generado: {date.today().strftime('%d/%m/%Y')}.")
    return _bytes(pdf)


def _pdf_requerimientos(reqs, nombre):
    pdf = _PDF("Reporte de Requerimientos", f"Evento: {nombre}  |  Total: {len(reqs)} items")
    pdf.add_page()
    s = _stats_numericos([float(r[3] or 0) for r in reqs])
    pdf.bloque_stats(s,"int")
    pdf.bloque_dist(_conteo([str(r[2]) for r in reqs]),"Distribucion por Tipo de Recurso")
    # Resumen por tipo
    tipo_sum: dict = {}
    for r in reqs: tipo_sum[str(r[2])]=tipo_sum.get(str(r[2]),0)+int(r[3] or 0)
    pdf.seccion("Cantidad Total por Tipo de Recurso")
    ws2=[70,40,40]; als2=["L","R","R"]; tot_g=sum(tipo_sum.values())
    pdf.th(["Tipo de Recurso","Cantidad Total","% del Total"],ws2)
    for i,(k,v) in enumerate(sorted(tipo_sum.items(),key=lambda x:-x[1])):
        pdf.tf([k,str(v),f"{v/tot_g*100:.1f}%"],ws2,i,als2)
    pdf.tt(["TOTAL",str(tot_g),"100.0%"],ws2,als2); pdf.ln(2)
    pdf.seccion("Listado de Requerimientos")
    ws3=[110,40,30]; als3=["L","C","R"]
    pdf.th(["Descripcion","Tipo","Cantidad"],ws3)
    for i,r in enumerate(reqs): pdf.tf([str(r[1])[:55],str(r[2]),str(r[3])],ws3,i,als3)
    pdf.tt(["TOTAL UNIDADES","",str(int(sum([float(r[3] or 0) for r in reqs])))],ws3,als3)
    pdf.sep(); pdf.nota(f"Evento: {nombre}. Generado: {date.today().strftime('%d/%m/%Y')}.")
    return _bytes(pdf)


def _pdf_cotizaciones(cots, nombre):
    pdf = _PDF("Reporte de Cotizaciones", f"Evento: {nombre}  |  Total: {len(cots)} cotizaciones")
    pdf.add_page()
    s = _stats_numericos([float(c[3] or 0) for c in cots])
    pdf.bloque_stats(s,"currency")
    pdf.bloque_dist(_conteo([str(c[4]) for c in cots]),"Distribucion por Estado")
    pdf.bloque_dist(_conteo([str(c[1]) for c in cots]),"Cotizaciones por Proveedor")
    # Monto por estado
    em: dict={}
    for c in cots: em[str(c[4])]=em.get(str(c[4]),0.0)+float(c[3] or 0)
    pdf.seccion("Monto Total por Estado")
    ws2=[50,50,50]; als2=["C","R","R"]; tot_m=sum(em.values())
    pdf.th(["Estado","Monto Total S/","% del Total"],ws2)
    for i,(k,v) in enumerate(sorted(em.items(),key=lambda x:-x[1])):
        pdf.tf([k,format_currency(v),f"{v/tot_m*100:.1f}%"],ws2,i,als2) if tot_m else None
    pdf.tt(["TOTAL",format_currency(tot_m),"100.0%"],ws2,als2); pdf.ln(2)
    pdf.seccion("Listado de Cotizaciones")
    ws3=[42,22,28,22,66]; als3=["L","C","R","C","L"]
    pdf.th(["Proveedor","Fecha","Monto S/","Estado","Descripcion"],ws3); tot=0.0
    for i,c in enumerate(cots):
        m=float(c[3] or 0); tot+=m
        pdf.tf([str(c[1])[:20],_fmt_date(c[2]),format_currency(m),str(c[4]),str(c[5] or "")[:32]],ws3,i,als3)
    pdf.tt(["","",format_currency(tot),"TOTAL",""],ws3,als3)
    pdf.sep(); pdf.nota(f"Evento: {nombre}. Generado: {date.today().strftime('%d/%m/%Y')}.")
    return _bytes(pdf)


def _pdf_asistente(info_ev, planes_ev, reqs_ev, cots_ev):
    nombre = info_ev[1] if info_ev else "Evento"
    pdf = _PDF("Reporte Integral del Evento", f"Evento: {nombre}")
    pdf.add_page()
    pdf.seccion("Informacion General del Evento")
    if info_ev:
        for i,(k,v) in enumerate([("Nombre",info_ev[1]),("Estado",info_ev[6]),
                                    ("Fecha",_fmt_date(info_ev[4])),
                                    ("Monto Estimado S/",format_currency(info_ev[5] or 0))]):
            pdf.kv(k,str(v),alt=(i%2==1))
    pdf.ln(2)
    # Stats planes
    s_p=_stats_numericos([float(p[2] or 0) for p in planes_ev])
    pdf.seccion("Estadisticas de Planes")
    for i,(k,v) in enumerate([
        ("Cantidad de planes",str(s_p["count"])),("Presupuesto total S/",format_currency(s_p["sum"])),
        ("Presupuesto promedio",format_currency(s_p["mean"])),("Presupuesto mediana",format_currency(s_p["median"])),
        ("Presupuesto minimo",format_currency(s_p["min"])),("Presupuesto maximo",format_currency(s_p["max"])),
        ("Presupuesto rango",format_currency(s_p["range"])),("Desv. estandar",format_currency(s_p["std"])),
        ("Coef. variacion",f"{s_p['cv']:.2f}%"),("Q1",format_currency(s_p["q1"])),
        ("Q3",format_currency(s_p["q3"])),("IQR",format_currency(s_p["iqr"])),
    ]): pdf.kv(k,v,alt=(i%2==1))
    pdf.bloque_dist(_conteo([str(p[3]) for p in planes_ev]),"Estados de Planes")
    # Stats reqs
    s_r=_stats_numericos([float(r[3] or 0) for r in reqs_ev])
    pdf.seccion("Estadisticas de Requerimientos")
    for i,(k,v) in enumerate([
        ("Total requerimientos",str(s_r["count"])),("Cantidad total",str(int(s_r["sum"]))),
        ("Cantidad promedio",f"{s_r['mean']:.2f}"),("Cantidad mediana",f"{s_r['median']:.2f}"),
        ("Cantidad minima",str(int(s_r["min"]))),("Cantidad maxima",str(int(s_r["max"]))),
        ("Rango",str(int(s_r["range"]))),("Desv. estandar",f"{s_r['std']:.2f}"),
        ("Coef. variacion",f"{s_r['cv']:.2f}%"),
    ]): pdf.kv(k,v,alt=(i%2==1))
    pdf.bloque_dist(_conteo([str(r[2]) for r in reqs_ev]),"Tipos de Requerimientos")
    # Stats cots
    s_c=_stats_numericos([float(c[3] or 0) for c in cots_ev])
    pdf.seccion("Estadisticas de Cotizaciones")
    for i,(k,v) in enumerate([
        ("Total cotizaciones",str(s_c["count"])),("Monto total S/",format_currency(s_c["sum"])),
        ("Monto promedio",format_currency(s_c["mean"])),("Monto mediana",format_currency(s_c["median"])),
        ("Monto minimo",format_currency(s_c["min"])),("Monto maximo",format_currency(s_c["max"])),
        ("Rango",format_currency(s_c["range"])),("Desv. estandar",format_currency(s_c["std"])),
        ("Coef. variacion",f"{s_c['cv']:.2f}%"),("Q1",format_currency(s_c["q1"])),
        ("Q3",format_currency(s_c["q3"])),("IQR",format_currency(s_c["iqr"])),
    ]): pdf.kv(k,v,alt=(i%2==1))
    pdf.bloque_dist(_conteo([str(c[4]) for c in cots_ev]),"Estados de Cotizaciones")
    pdf.sep(); pdf.nota(f"Generado: {date.today().strftime('%d/%m/%Y')}. Reporte de caracter informativo.")
    return _bytes(pdf)


# ══════════════════════════════════════════════════════════════════
#  DIALOGS
# ══════════════════════════════════════════════════════════════════

ESTADOS_PLAN = ['Borrador', 'En Revisión', 'Aprobado', 'Rechazado', 'Registrado']


@st.dialog("Elaborar Nuevo Plan")
def _dialog_elaborar_plan(id_evento: int, nombre_evento: str):
    st.markdown(f"**Evento:** {nombre_evento}")
    fecha_elab  = st.date_input("Fecha de elaboración", value=date.today(), format="DD/MM/YYYY")
    presupuesto = st.number_input("Presupuesto S/", min_value=0.0, step=100.0)
    descripcion = st.text_area("Descripción del plan")
    b1, b2 = st.columns(2)
    if b1.button("Guardar Plan", type="primary"):
        if model_plan_evento.create(id_evento, fecha_elab, presupuesto, descripcion):
            st.success("Plan creado."); st.rerun()
    if b2.button("Cancelar"): st.rerun()


@st.dialog("Editar Plan")
def _dialog_editar_plan(id_plan, id_evento, nombre_evento, fecha_actual, presupuesto_actual, estado_actual, descripcion_actual):
    st.markdown(f"**Evento:** {nombre_evento}")
    fecha_elab  = st.date_input("Fecha", value=fecha_actual or date.today(), format="DD/MM/YYYY", key=f"pe_f_{id_plan}")
    c1, c2      = st.columns(2)
    presupuesto = c1.number_input("Presupuesto S/", min_value=0.0, step=100.0, value=float(presupuesto_actual or 0), key=f"pe_p_{id_plan}")
    idx_est     = ESTADOS_PLAN.index(estado_actual) if estado_actual in ESTADOS_PLAN else 0
    estado      = c2.selectbox("Estado", ESTADOS_PLAN, index=idx_est, key=f"pe_e_{id_plan}")
    descripcion = st.text_area("Descripción", value=descripcion_actual or "", key=f"pe_d_{id_plan}")
    b1, b2 = st.columns(2)
    if b1.button("Guardar", key=f"pe_s_{id_plan}", type="primary"):
        ok_u = model_plan_evento.update(id_plan, fecha_elab, presupuesto, descripcion)
        ok_e = True
        if estado != estado_actual:
            ok_e = model_plan_evento.cambiar_estado(id_plan, estado)
            if estado == 'Aprobado':   model_evento.cambiar_estado(id_evento, 'Plan Aprobado')
            if estado == 'Registrado': model_evento.cambiar_estado(id_evento, 'Confirmada')
        if ok_u and ok_e:
            st.success("Actualizado."); st.rerun()
    if b2.button("Cerrar", key=f"pe_c_{id_plan}"): st.rerun()


@st.dialog("Agregar Requerimiento")
def _dialog_agregar_req(id_evento: int, nombre_evento: str):
    st.markdown(f"**Evento:** {nombre_evento}")
    d1, d2 = st.columns(2)
    desc_r = d1.text_input("Descripción *", key=f"rn_d_{id_evento}")
    tipo_r = d2.selectbox("Tipo de recurso", model_requerimiento.TIPOS_RECURSO, key=f"rn_t_{id_evento}")
    cant_r = st.number_input("Cantidad", min_value=1, step=1, key=f"rn_c_{id_evento}")
    b1, b2 = st.columns(2)
    if b1.button("Guardar", key=f"rn_s_{id_evento}", type="primary"):
        if not desc_r: st.error("La descripción es obligatoria.")
        elif model_requerimiento.create(id_evento, desc_r, tipo_r, int(cant_r)):
            st.success("Requerimiento agregado."); st.rerun()
    if b2.button("Cancelar", key=f"rn_x_{id_evento}"): st.rerun()


@st.dialog("Editar Requerimiento")
def _dialog_editar_req(id_req, desc_act, tipo_act, cant_act):
    e1, e2   = st.columns(2)
    desc_new = e1.text_input("Descripción", value=desc_act or "", key=f"re_d_{id_req}")
    idx_t    = model_requerimiento.TIPOS_RECURSO.index(tipo_act) if tipo_act in model_requerimiento.TIPOS_RECURSO else 0
    tipo_new = e2.selectbox("Tipo", model_requerimiento.TIPOS_RECURSO, index=idx_t, key=f"re_t_{id_req}")
    cant_new = st.number_input("Cantidad", min_value=1, step=1, value=int(cant_act or 1), key=f"re_c_{id_req}")
    b1, b2   = st.columns(2)
    if b1.button("Guardar", key=f"re_s_{id_req}", type="primary"):
        if not desc_new: st.error("La descripción es obligatoria.")
        elif model_requerimiento.update(id_req, desc_new, tipo_new, int(cant_new)):
            st.success("Actualizado."); st.rerun()
    if b2.button("Cerrar", key=f"re_x_{id_req}"): st.rerun()

@st.dialog("Editar Evento")
def _dialog_editar_evento(id_evento, nombre_act, tipo_act, lugar_act, fecha_act, monto_act, estado_act):
    from config import ESTADOS_EVENTO
    st.markdown(f"**ID:** {id_evento}")
    f1, f2 = st.columns(2)
    nombre = f1.text_input("Nombre", value=str(nombre_act or ""), key=f"evd_n_{id_evento}")
    tipos = ["Corporativo","Social","Institucional","Cultural","Deportivo","Otro"]
    idx_t = tipos.index(tipo_act) if tipo_act in tipos else 0
    tipo = f2.selectbox("Tipo", tipos, index=idx_t, key=f"evd_t_{id_evento}")
    lugar = f1.text_input("Lugar", value=str(lugar_act or ""), key=f"evd_l_{id_evento}")
    fecha = f2.date_input("Fecha", value=fecha_act or date.today(), key=f"evd_f_{id_evento}", format="DD/MM/YYYY")
    monto = st.number_input("Monto S/", min_value=0.0, value=float(monto_act or 0), key=f"evd_m_{id_evento}")
    idx_e = ESTADOS_EVENTO.index(estado_act) if estado_act in ESTADOS_EVENTO else 0
    estado = st.selectbox("Estado", ESTADOS_EVENTO, index=idx_e, key=f"evd_e_{id_evento}")
    b1, b2 = st.columns(2)
    if b1.button("Guardar", key=f"evd_s_{id_evento}", type="primary"):
        ok_u = model_evento.update(id_evento, nombre, tipo, lugar, fecha, monto)
        ok_e = model_evento.cambiar_estado(id_evento, estado) if estado != estado_act else True
        if ok_u and ok_e:
            st.success("Evento actualizado."); st.rerun()
    if b2.button("Cerrar", key=f"evd_x_{id_evento}"): st.rerun()


@st.dialog("Registrar Cotización")
def _dialog_registrar_cot(id_evento: int, nombre_evento: str):
    st.markdown(f"**Evento:** {nombre_evento}")
    proveedores = model_proveedor.get_all()
    if not proveedores:
        st.info("No hay proveedores registrados.")
        if st.button("Cerrar", key=f"cn_x_{id_evento}"): st.rerun()
        return
    prov_sel = st.selectbox("Proveedor *", options=proveedores, format_func=lambda p: p[1], key=f"cn_p_{id_evento}")
    fecha_c  = st.date_input("Fecha", value=date.today(), format="DD/MM/YYYY", key=f"cn_f_{id_evento}")
    monto_c  = st.number_input("Monto S/", min_value=0.0, key=f"cn_m_{id_evento}")
    desc_c   = st.text_area("Descripción", key=f"cn_d_{id_evento}")
    b1, b2   = st.columns(2)
    if b1.button("Guardar", key=f"cn_s_{id_evento}", type="primary"):
        if model_cotizacion.create(int(prov_sel[0]), int(id_evento), fecha_c, monto_c, desc_c):
            st.success("Cotización registrada."); st.rerun()
    if b2.button("Cancelar", key=f"cn_x2_{id_evento}"): st.rerun()


@st.dialog("Editar Cotización")
def _dialog_editar_cot(id_cot, proveedor, fecha_actual, monto_actual, estado_actual, desc_actual):
    st.markdown(f"**Proveedor:** {proveedor}  |  **Fecha:** {_fmt_date(fecha_actual)}")
    c1, c2   = st.columns(2)
    monto_c  = c1.number_input("Monto S/", min_value=0.0, value=float(monto_actual or 0), key=f"ce_m_{id_cot}")
    estados  = ["Pendiente", "Aceptada", "Rechazada"]
    idx_est  = estados.index(estado_actual) if estado_actual in estados else 0
    estado_c = c2.selectbox("Estado", estados, index=idx_est, key=f"ce_e_{id_cot}")
    desc_c   = st.text_area("Descripción", value=desc_actual or "", key=f"ce_d_{id_cot}")
    b1, b2   = st.columns(2)
    if b1.button("Guardar", key=f"ce_s_{id_cot}", type="primary"):
        ok_u = model_cotizacion.update(int(id_cot), monto_c, desc_c)
        ok_e = model_cotizacion.cambiar_estado(int(id_cot), estado_c) if estado_c != estado_actual else True
        if ok_u and ok_e:
            st.success("Actualizada."); st.rerun()
    if b2.button("Cerrar", key=f"ce_x_{id_cot}"): st.rerun()


# ══════════════════════════════════════════════════════════════════
#  PAGINADOR
# ══════════════════════════════════════════════════════════════════

def _paginar(items, size, pk, sk=None, snow=None):
    total = max(1, (len(items) + size - 1) // size)
    if sk and snow is not None:
        if st.session_state.get(sk) != snow:
            st.session_state[sk] = snow; st.session_state[pk] = 1
    if st.session_state.get(pk, 1) > total: st.session_state[pk] = total
    page  = int(st.session_state.get(pk, 1))
    start = (page - 1) * size; end = start + size

    def _nav():
        _, nc, _ = st.columns([2, 0.7, 2])
        with nc:
            p1, p2, p3 = st.columns([0.6, 0.6, 0.6], vertical_alignment="center")
            if p1.button("⏮️", key=f"{pk}_p", disabled=(page <= 1)):
                st.session_state[pk] = max(1, page - 1); st.rerun()
            p2.markdown(f"**{page}/{total}**")
            if p3.button("⏭️", key=f"{pk}_n", disabled=(page >= total)):
                st.session_state[pk] = min(total, page + 1); st.rerun()

    return items[start:end], start, end, _nav


# ══════════════════════════════════════════════════════════════════
#  SHOW — PUNTO DE ENTRADA
# ══════════════════════════════════════════════════════════════════

def show():
    st.markdown(THEME_CSS, unsafe_allow_html=True)
    st.markdown('<div id="planificacion-page">', unsafe_allow_html=True)

    requiere_rol(['Administrador', 'Jefe de Planificación', 'Jefe de Eventos'])
    st.title("📋 Gestión de Planificación de Eventos")

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📅 Eventos", "📝 Plan del Evento",
        "📦 Requerimientos", "💰 Cotizaciones", "🧭 Asistente",
    ])

    # ════════════════════════════════════════════════════════════════
    # TAB 1 — EVENTOS
    # ════════════════════════════════════════════════════════════════
    with tab1:
        with st.expander("➕ Registrar Nuevo Evento", expanded=False):
            from cu1_contratos import model_cliente
            clientes_rows = model_cliente.get_activos()
            with st.form("form_evento"):
                c1, c2      = st.columns(2)
                nombre       = c1.text_input("Nombre del evento *")
                tipo_evento  = c2.selectbox("Tipo", ["Corporativo","Social","Institucional","Cultural","Deportivo","Otro"])
                lugar_evento = c1.text_input("Lugar")
                fecha_evento = c2.date_input("Fecha del evento", format="DD/MM/YYYY")
                monto_evento = st.number_input("Monto estimado S/", min_value=0.0)
                cliente_sel  = st.selectbox("Cliente *",
                    options=clientes_rows if clientes_rows else [None],
                    format_func=lambda c: c[1] if c else "Sin clientes")
                if st.form_submit_button("Registrar Evento"):
                    if not nombre or not clientes_rows:
                        st.error("Nombre y cliente son obligatorios.")
                    elif model_evento.create(nombre, tipo_evento, lugar_evento, fecha_evento, monto_evento, cliente_sel[0]):
                        st.success(f"Evento '{nombre}' registrado."); st.rerun()

        st.divider()
        st.subheader("Listado de Eventos")
        rows = model_evento.get_all()
        if not rows:
            st.info("No hay eventos registrados.")
        else:
            from config import ESTADOS_EVENTO
            f1,f2,f3,f4,f5 = st.columns([2,1.4,1.4,2.2,2.5])
            q_nom   = f1.text_input("Filtrar por nombre", key="ev_q_n")
            q_tipo  = f2.selectbox("Tipo", ["Todos","Corporativo","Social","Institucional","Cultural","Deportivo","Otro"], key="ev_q_t")
            q_est   = f3.selectbox("Estado", ["Todos"]+ESTADOS_EVENTO, key="ev_q_e")
            min_f   = min((r[4] for r in rows if r[4]), default=date.today())
            q_fdes  = f4.date_input("A partir de", value=min_f, key="ev_q_fd", format="DD/MM/YYYY")
            precios = [float(r[5] or 0) for r in rows]
            mn_p, mx_p = min(precios,default=0.0), max(precios,default=10000.0)
            if mn_p == mx_p: mx_p = mn_p + 1000.0
            q_prec  = f5.slider("Rango S/", float(mn_p), float(mx_p), (float(mn_p),float(mx_p)), key="ev_q_pr")

            filtered = [r for r in rows
                if (not q_nom or q_nom.lower() in str(r[1]).lower())
                and (q_tipo=="Todos" or r[2]==q_tipo)
                and (q_est=="Todos" or r[6]==q_est)
                and q_prec[0]<=float(r[5] or 0)<=q_prec[1]
                and (not r[4] or r[4]>=q_fdes)]

            st.markdown(f"**{len(filtered)}** eventos encontrados")
            paged, start, end, nav = _paginar(filtered, 5, "ev_pg",
                "ev_fsig", (q_nom,q_tipo,q_est,str(q_fdes),q_prec[0],q_prec[1]))
            st.caption(f"Mostrando {start+1}–{min(end,len(filtered))} de {len(filtered)}")

            for lbl,col in zip(["Nombre","Tipo","Lugar","Fecha","Monto","Estado","Acc."],
                                 st.columns([3,2,3,2,2,2,1.6])):
                col.markdown(f"**{lbl}**")
            for r in paged:
                cols = st.columns([3,2,3,2,2,2,1.6])
                cols[0].write(str(r[1])); cols[1].write(str(r[2]))
                cols[2].write(str(r[3] or "")); cols[3].write(_fmt_date(r[4]))
                cols[4].write(format_currency(r[5] or 0)); cols[5].write(str(r[6]))
                a1,a2 = cols[6].columns(2)
                if a1.button("✏️", key=f"ev_ed_{r[0]}", help="Editar"):
                    _dialog_editar_evento(r[0], r[1], r[2], r[3], r[4], r[5], r[6])
                if a2.button("🗑️", key=f"ev_dl_{r[0]}", help="Eliminar"):
                    if model_evento.delete(r[0]): st.success("Eliminado."); st.rerun()
            nav()

            with st.expander("📊 Estadísticas y Gráficos", expanded=False):
                montos = [float(r[5] or 0) for r in filtered]
                s = _stats_numericos(montos)
                st.markdown("#### 💰 Estadísticas de Montos")
                _panel_stats(s, "currency")
                st.divider()
                g1,g2 = st.columns(2)
                est_cnt  = _conteo([r[6] for r in filtered])
                tipo_cnt = _conteo([r[2] for r in filtered])
                with g1:
                    _show(_fig_dona(est_cnt, "Distribución por Estado"))
                    _show(_fig_barras_h(tipo_cnt, "Eventos por Tipo", "Cantidad"))
                with g2:
                    _show(_fig_dona(tipo_cnt, "Distribución por Tipo"))
                    _show(_fig_histo(montos, "Distribución de Montos", " S/"))
                fechas_ev=[r[4] for r in filtered if r[4]]
                montos_ev=[float(r[5] or 0) for r in filtered if r[4]]
                _show(_fig_linea(fechas_ev, montos_ev, "Monto de Eventos por Fecha", " S/"))
                st.divider()
                st.download_button("📥 Exportar Reporte PDF", data=_pdf_eventos(filtered),
                    file_name="reporte_eventos.pdf", mime="application/pdf", use_container_width=True)

    # ════════════════════════════════════════════════════════════════
    # TAB 2 — PLAN DEL EVENTO
    # ════════════════════════════════════════════════════════════════
    with tab2:
        ev_act = model_evento.get_activos()
        if not ev_act: st.warning("No hay eventos activos.")
        else:
            st.subheader("Elaborar Nuevo Plan")
            ev_b  = st.text_input("Buscar evento", key="pl_ev_b")
            ev_f  = [e for e in ev_act if ev_b.strip().lower() in str(e[1]).lower()] if ev_b else ev_act
            id_ev_sel=None; nom_ev_sel=None
            if not ev_f: st.info("No se encontraron eventos.")
            else:
                lc, rc = st.columns([4,1], vertical_alignment="bottom")
                ev_s  = lc.selectbox("Evento", ev_f, format_func=lambda e:e[1], key="pl_ev_s", label_visibility="collapsed")
                id_ev_sel=int(ev_s[0]); nom_ev_sel=str(ev_s[1])
                if rc.button("➕ Elaborar", use_container_width=True):
                    _dialog_elaborar_plan(id_ev_sel, nom_ev_sel)

            st.divider(); st.subheader("Planes registrados")
            if not id_ev_sel: st.info("Seleccione un evento.")
            else:
                planes = model_plan_evento.get_by_evento(id_ev_sel)
                if not planes: st.info("No hay planes para este evento.")
                else:
                    paged,start,end,nav = _paginar(planes,5,f"pl_pg_{id_ev_sel}")
                    st.caption(f"{start+1}–{min(end,len(planes))} de {len(planes)}")
                    for lbl,col in zip(["Fecha Elab.","Presupuesto","Estado","Descripción","Acc."],
                                         st.columns([2,2,2,4,1.4])):
                        col.markdown(f"**{lbl}**")
                    for row in paged:
                        c=st.columns([2,2,2,4,1.4])
                        c[0].write(_fmt_date(row[1])); c[1].write(format_currency(row[2] or 0))
                        c[2].write(str(row[3])); dtxt=str(row[4] or "")
                        c[3].write(dtxt[:80]+("…" if len(dtxt)>80 else ""))
                        a1,a2=c[4].columns(2)
                        if a1.button("✏️", key=f"pl_ed_{row[0]}", help="Editar"):
                            _dialog_editar_plan(int(row[0]),id_ev_sel,nom_ev_sel,row[1],float(row[2] or 0),str(row[3]),dtxt)
                        if a2.button("🗑️", key=f"pl_dl_{row[0]}", help="Eliminar"):
                            if model_plan_evento.delete(int(row[0])): st.success("Eliminado."); st.rerun()
                    nav()

                    with st.expander("📊 Estadísticas y Gráficos", expanded=False):
                        pres = [float(p[2] or 0) for p in planes]
                        s = _stats_numericos(pres)
                        st.markdown("#### 💰 Estadísticas de Presupuestos")
                        _panel_stats(s,"currency")
                        st.divider()
                        g1,g2=st.columns(2)
                        with g1: _show(_fig_dona(_conteo([str(p[3]) for p in planes]),"Estados del Plan"))
                        with g2: _show(_fig_histo(pres,"Distribución de Presupuestos"," S/"))
                        _show(_fig_linea([p[1] for p in planes if p[1]],
                                         [float(p[2] or 0) for p in planes if p[1]],
                                         "Evolución de Presupuestos"," S/"))
                        st.divider()
                        st.download_button("📥 Exportar Reporte PDF",
                            data=_pdf_planes(planes,nom_ev_sel),
                            file_name=f"reporte_planes_{id_ev_sel}.pdf",
                            mime="application/pdf", use_container_width=True)

    # ════════════════════════════════════════════════════════════════
    # TAB 3 — REQUERIMIENTOS
    # ════════════════════════════════════════════════════════════════
    with tab3:
        ev_r = model_evento.get_activos()
        if not ev_r: st.warning("No hay eventos activos.")
        else:
            st.subheader("Agregar Requerimiento")
            ev_br  = st.text_input("Buscar evento", key="rq_ev_b")
            ev_fr  = [e for e in ev_r if ev_br.strip().lower() in str(e[1]).lower()] if ev_br else ev_r
            id_ev_r=None; nom_ev_r=None
            if not ev_fr: st.info("No se encontraron eventos.")
            else:
                lc,rc=st.columns([4,1],vertical_alignment="bottom")
                ev_sr=lc.selectbox("Evento",ev_fr,format_func=lambda e:e[1],key="rq_ev_s",label_visibility="collapsed")
                id_ev_r=int(ev_sr[0]); nom_ev_r=str(ev_sr[1])
                if rc.button("➕ Agregar", use_container_width=True): _dialog_agregar_req(id_ev_r,nom_ev_r)

            st.divider()
            t1,t2=st.columns([2,3],vertical_alignment="bottom"); t1.subheader("Requerimientos")
            if not id_ev_r: st.info("Seleccione un evento.")
            else:
                reqs=model_requerimiento.get_by_evento(id_ev_r)
                if not reqs: st.info("No hay requerimientos.")
                else:
                    fd,ft=t2.columns([2,1.3],vertical_alignment="bottom")
                    q_desc=fd.text_input("Buscar",key="rq_fd",label_visibility="collapsed",placeholder="Buscar descripción")
                    q_tipo=ft.selectbox("Tipo",["Todos"]+list(model_requerimiento.TIPOS_RECURSO),key="rq_ft",label_visibility="collapsed")
                    rf=[r for r in reqs
                        if (not q_desc or q_desc.strip().lower() in str(r[1]).lower())
                        and (q_tipo=="Todos" or str(r[2])==q_tipo)]
                    paged,start,end,nav=_paginar(rf,5,f"rq_pg_{id_ev_r}",f"rq_fs_{id_ev_r}",(q_desc or "",q_tipo or "Todos"))
                    st.caption(f"{start+1}–{min(end,len(rf))} de {len(rf)}")
                    for lbl,col in zip(["Descripción","Tipo","Cantidad","Acc."],st.columns([5,2,1.5,1.4])):
                        col.markdown(f"**{lbl}**")
                    for req in paged:
                        rid,rdesc,rtipo,rcant=req
                        c=st.columns([5,2,1.5,1.4])
                        c[0].write(str(rdesc)); c[1].write(str(rtipo)); c[2].write(str(rcant))
                        a1,a2=c[3].columns(2)
                        if a1.button("✏️",key=f"rq_ed_{rid}",help="Editar"):
                            _dialog_editar_req(int(rid),str(rdesc or ""),str(rtipo),int(rcant))
                        if a2.button("🗑️",key=f"rq_dl_{rid}",help="Eliminar"):
                            if model_requerimiento.delete(int(rid)): st.success("Eliminado."); st.rerun()
                    nav()

                    with st.expander("📊 Estadísticas y Gráficos", expanded=False):
                        cants=[float(r[3] or 0) for r in rf]
                        s=_stats_numericos(cants)
                        st.markdown("#### 📦 Estadísticas de Cantidades")
                        _panel_stats(s,"int"); st.divider()
                        tipo_cnt=_conteo([str(r[2]) for r in rf])
                        tipo_sum={k:sum(int(r[3] or 0) for r in rf if str(r[2])==k) for k in tipo_cnt}
                        g1,g2=st.columns(2)
                        with g1: _show(_fig_dona(tipo_cnt,"Requerimientos por Tipo"))
                        with g2: _show(_fig_barras_h(tipo_sum,"Cantidad Total por Tipo","Unidades"))
                        _show(_fig_histo(cants,"Distribución de Cantidades"))
                        st.divider()
                        st.download_button("📥 Exportar Reporte PDF",
                            data=_pdf_requerimientos(rf,nom_ev_r),
                            file_name=f"reporte_reqs_{id_ev_r}.pdf",
                            mime="application/pdf", use_container_width=True)

                    st.divider()
                    with st.expander("🔍 Verificar Disponibilidad Interna", expanded=False):
                        for req in rf:
                            disponibles=model_recurso.get_disponibles_por_tipo(req[2])
                            total_disp=sum(int(r[3]) for r in disponibles)
                            if total_disp>=req[3]:
                                st.success(f"✅ {req[1]}: {total_disp} disponibles (necesita {req[3]})")
                            else:
                                st.warning(f"⚠️ {req[1]}: solo {total_disp} disponibles (necesita {req[3]}) — considerar cotización")

    # ════════════════════════════════════════════════════════════════
    # TAB 4 — COTIZACIONES
    # ════════════════════════════════════════════════════════════════
    with tab4:
        ev_c=model_evento.get_activos()
        if not ev_c: st.warning("No hay eventos activos.")
        else:
            st.subheader("Registrar Cotización")
            ev_bc=st.text_input("Buscar evento",key="ct_ev_b")
            ev_fc=[e for e in ev_c if ev_bc.strip().lower() in str(e[1]).lower()] if ev_bc else ev_c
            id_ev_c=None; nom_ev_c=None
            if not ev_fc: st.info("No se encontraron eventos.")
            else:
                lc,rc=st.columns([4,1],vertical_alignment="bottom")
                ev_sc=lc.selectbox("Evento",ev_fc,format_func=lambda e:e[1],key="ct_ev_s",label_visibility="collapsed")
                id_ev_c=int(ev_sc[0]); nom_ev_c=str(ev_sc[1])
                if rc.button("➕ Registrar",use_container_width=True): _dialog_registrar_cot(id_ev_c,nom_ev_c)

            st.divider()
            t1,t2=st.columns([2,3],vertical_alignment="bottom"); t1.subheader("Cotizaciones")
            if not id_ev_c: st.info("Seleccione un evento.")
            else:
                cots=model_cotizacion.get_by_evento(id_ev_c)
                if not cots: st.info("No hay cotizaciones.")
                else:
                    fp,fe=t2.columns([2,1.3],vertical_alignment="bottom")
                    q_prov=fp.text_input("Buscar proveedor",key="ct_fp",label_visibility="collapsed",placeholder="Buscar proveedor")
                    q_est=fe.selectbox("Estado",["Todos","Pendiente","Aceptada","Rechazada"],key="ct_fe",label_visibility="collapsed")
                    cf=[c for c in cots
                        if (not q_prov or q_prov.strip().lower() in str(c[1]).lower())
                        and (q_est=="Todos" or str(c[4])==q_est)]
                    paged,start,end,nav=_paginar(cf,5,f"ct_pg_{id_ev_c}",f"ct_fs_{id_ev_c}",(q_prov or "",q_est or "Todos"))
                    st.caption(f"{start+1}–{min(end,len(cf))} de {len(cf)}")
                    for lbl,col in zip(["Proveedor","Fecha","Monto","Estado","Descripción","Acciones"],
                                         st.columns([3,1.6,1.6,1.6,4,4.6])):
                        col.markdown(f"**{lbl}**")
                    for cot in paged:
                        cid,prov,fecha,monto,estado,desc=cot
                        c=st.columns([3,1.6,1.6,1.6,4,4.6])
                        c[0].write(str(prov)); c[1].write(_fmt_date(fecha))
                        c[2].write(format_currency(monto or 0)); c[3].write(str(estado))
                        dtxt=str(desc or ""); c[4].write(dtxt[:80]+("…" if len(dtxt)>80 else ""))
                        _,a1,a2,a3,a4,_=c[5].columns([0.6,1,1,1,1,0.6],gap="small",vertical_alignment="center")
                        if a1.button("✏️",key=f"ct_ed_{cid}",help="Editar"):
                            _dialog_editar_cot(int(cid),str(prov),fecha,float(monto or 0),str(estado),dtxt)
                        if a2.button("🗑️",key=f"ct_dl_{cid}",help="Eliminar"):
                            if model_cotizacion.delete(int(cid)): st.success("Eliminada."); st.rerun()
                            else: st.info("Solo se puede eliminar si está Pendiente.")
                        if a3.button("✅",key=f"ct_ac_{cid}",help="Aceptar"):
                            if model_cotizacion.cambiar_estado(int(cid),"Aceptada"): st.success("Aceptada."); st.rerun()
                        if a4.button("❌",key=f"ct_rj_{cid}",help="Rechazar"):
                            if model_cotizacion.cambiar_estado(int(cid),"Rechazada"): st.warning("Rechazada."); st.rerun()
                    nav()

                    with st.expander("📊 Estadísticas y Gráficos", expanded=False):
                        montos=[float(c[3] or 0) for c in cf]
                        s=_stats_numericos(montos)
                        st.markdown("#### 💰 Estadísticas de Montos")
                        _panel_stats(s,"currency"); st.divider()
                        est_cnt=_conteo([str(c[4]) for c in cf])
                        prov_cnt=_conteo([str(c[1]) for c in cf])
                        g1,g2=st.columns(2)
                        with g1:
                            _show(_fig_dona(est_cnt,"Distribución por Estado"))
                            _show(_fig_histo(montos,"Distribución de Montos"," S/"))
                        with g2:
                            _show(_fig_barras_h(prov_cnt,"Cotizaciones por Proveedor","Cant."))
                        # Monto por estado
                        m_est={e:sum(float(c[3] or 0) for c in cf if str(c[4])==e)
                               for e in ["Aceptada","Pendiente","Rechazada"]}
                        _show(_fig_barras_h({k:v for k,v in m_est.items() if v>0},
                                            "Monto Total por Estado","S/"))
                        st.divider()
                        st.download_button("📥 Exportar Reporte PDF",
                            data=_pdf_cotizaciones(cf,nom_ev_c),
                            file_name=f"reporte_cot_{id_ev_c}.pdf",
                            mime="application/pdf", use_container_width=True)

    # ════════════════════════════════════════════════════════════════
    # TAB 5 — ASISTENTE
    # ════════════════════════════════════════════════════════════════
    with tab5:
        eventos=model_evento.get_activos()
        if not eventos: st.warning("No hay eventos activos.")
        else:
            st.subheader("Asistente de Planificación")
            ev_ba=st.text_input("Buscar evento",key="as_ev_b")
            ev_fa=[e for e in eventos if ev_ba.strip().lower() in str(e[1]).lower()] if ev_ba else eventos
            if not ev_fa: st.info("No se encontraron eventos."); return

            ev_sel=st.selectbox("Evento",ev_fa,format_func=lambda e:e[1],key="as_ev_s",label_visibility="collapsed")
            id_ev=int(ev_sel[0])
            info_ev   = model_evento.get_by_id(id_ev)
            planes_ev = model_plan_evento.get_by_evento(id_ev)
            reqs_ev   = model_requerimiento.get_by_evento(id_ev)
            cots_ev   = model_cotizacion.get_by_evento(id_ev)

            st.subheader("Resumen General")
            ca,cb=st.columns(2)
            with ca:
                if info_ev:
                    st.metric("Evento",info_ev[1]); st.metric("Estado",info_ev[6])
                    st.metric("Fecha",_fmt_date(info_ev[4])); st.metric("Monto",format_currency(info_ev[5] or 0))
            with cb:
                tot_pres=sum(float(p[2] or 0) for p in planes_ev)
                tot_cot =sum(float(c[3] or 0) for c in cots_ev)
                st.metric("Planes",len(planes_ev)); st.metric("Requerimientos",len(reqs_ev))
                st.metric("Cotizaciones",len(cots_ev))
                st.metric("Presupuesto Total",format_currency(tot_pres))
                st.metric("Total Cotizado",format_currency(tot_cot))

            st.divider()
            with st.expander("📊 Estadísticas Descriptivas Completas", expanded=False):
                s_p=_stats_numericos([float(p[2] or 0) for p in planes_ev])
                s_r=_stats_numericos([float(r[3] or 0) for r in reqs_ev])
                s_c=_stats_numericos([float(c[3] or 0) for c in cots_ev])

                st.markdown("#### 📝 Presupuestos (Planes)")
                _panel_stats(s_p,"currency")
                st.markdown("#### 📦 Cantidades (Requerimientos)")
                _panel_stats(s_r,"int")
                st.markdown("#### 💰 Montos (Cotizaciones)")
                _panel_stats(s_c,"currency")

                st.divider()
                g1,g2,g3=st.columns(3)
                with g1: _show(_fig_dona(_conteo([str(p[3]) for p in planes_ev]),"Estados Planes"))
                with g2: _show(_fig_dona(_conteo([str(r[2]) for r in reqs_ev]),"Tipos Req."))
                with g3: _show(_fig_dona(_conteo([str(c[4]) for c in cots_ev]),"Estados Cot."))

                _show(_fig_linea([c[2] for c in cots_ev if c[2]],
                                  [float(c[3] or 0) for c in cots_ev if c[2]],
                                  "Evolución de Cotizaciones"," S/"))
                _show(_fig_barras_h(
                    {"Presupuesto Total":tot_pres,"Total Cotizado":tot_cot,
                     "Monto Evento":float(info_ev[5] or 0) if info_ev else 0},
                    "Comparativa de Montos (S/)","S/"))

                st.divider()
                st.download_button("📥 Exportar Reporte Integral PDF",
                    data=_pdf_asistente(info_ev,planes_ev,reqs_ev,cots_ev),
                    file_name=f"reporte_integral_{id_ev}.pdf",
                    mime="application/pdf", use_container_width=True)

            st.subheader("Listados")
            with st.expander("Planes",expanded=False):
                if not planes_ev: st.info("Sin planes.")
                else:
                    paged,start,end,nav=_paginar(planes_ev,5,f"as_pl_{id_ev}")
                    st.caption(f"{start+1}–{min(end,len(planes_ev))} de {len(planes_ev)}")
                    for lbl,col in zip(["Fecha","Presupuesto","Estado","Descripción"],st.columns([2,2,2,4])):
                        col.markdown(f"**{lbl}**")
                    for row in paged:
                        c=st.columns([2,2,2,4]); c[0].write(_fmt_date(row[1]))
                        c[1].write(format_currency(row[2] or 0)); c[2].write(str(row[3]))
                        dtxt=str(row[4] or ""); c[3].write(dtxt[:120]+("…" if len(dtxt)>120 else ""))
                    nav()

            with st.expander("Requerimientos",expanded=False):
                if not reqs_ev: st.info("Sin requerimientos.")
                else:
                    paged,start,end,nav=_paginar(reqs_ev,5,f"as_rq_{id_ev}")
                    st.caption(f"{start+1}–{min(end,len(reqs_ev))} de {len(reqs_ev)}")
                    for lbl,col in zip(["Descripción","Tipo","Cantidad"],st.columns([5,2,1.5])):
                        col.markdown(f"**{lbl}**")
                    for row in paged:
                        c=st.columns([5,2,1.5]); c[0].write(str(row[1])); c[1].write(str(row[2])); c[2].write(str(row[3]))
                    nav()

            with st.expander("Cotizaciones",expanded=False):
                if not cots_ev: st.info("Sin cotizaciones.")
                else:
                    paged,start,end,nav=_paginar(cots_ev,5,f"as_ct_{id_ev}")
                    st.caption(f"{start+1}–{min(end,len(cots_ev))} de {len(cots_ev)}")
                    for lbl,col in zip(["Proveedor","Fecha","Monto","Estado","Descripción"],st.columns([3,1.6,1.6,1.6,4])):
                        col.markdown(f"**{lbl}**")
                    for row in paged:
                        c=st.columns([3,1.6,1.6,1.6,4]); c[0].write(str(row[1])); c[1].write(_fmt_date(row[2]))
                        c[2].write(format_currency(row[3] or 0)); c[3].write(str(row[4]))
                        dtxt=str(row[5] or ""); c[4].write(dtxt[:120]+("…" if len(dtxt)>120 else ""))
                    nav()

            st.divider()
            # ── Export Word ──────────────────────────────────────────
            doc=Document(); doc.add_heading("Planificacion de Evento",0)
            if info_ev:
                for lbl,val in [("Evento",info_ev[1]),("Estado",info_ev[6]),
                                  ("Fecha",_fmt_date(info_ev[4])),("Monto",format_currency(info_ev[5] or 0))]:
                    p=doc.add_paragraph(); p.add_run(f"{lbl}: ").bold=True; p.add_run(str(val))
            for heading,data,headers in [
                ("Planes",planes_ev,["ID","Fecha","Presupuesto","Estado","Descripcion"]),
                ("Requerimientos",reqs_ev,["ID","Descripcion","Tipo","Cantidad"]),
                ("Cotizaciones",cots_ev,["ID","Proveedor","Fecha","Monto","Estado","Descripcion"]),
            ]:
                doc.add_heading(heading,level=1)
                if data:
                    tbl=doc.add_table(rows=1,cols=len(headers))
                    for i,h in enumerate(headers): tbl.rows[0].cells[i].text=h
                    for row in data:
                        cells=tbl.add_row().cells
                        if heading=="Planes":
                            cells[0].text=str(row[0]); cells[1].text=_fmt_date(row[1])
                            cells[2].text=format_currency(row[2] or 0); cells[3].text=str(row[3]); cells[4].text=str(row[4] or "")
                        elif heading=="Requerimientos":
                            for j,v in enumerate(row): cells[j].text=str(v)
                        else:
                            cells[0].text=str(row[0]); cells[1].text=str(row[1]); cells[2].text=_fmt_date(row[2])
                            cells[3].text=format_currency(row[3] or 0); cells[4].text=str(row[4]); cells[5].text=str(row[5] or "")
            buf=BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button("📥 Exportar a Word", data=buf,
                file_name=f"Planificacion_Evento_{id_ev}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)
